import os
import math
from openai import OpenAI
from fpdf import FPDF
from dotenv import load_dotenv
from pydub import AudioSegment
import tempfile
import tiktoken
import subprocess
import shutil
from docx import Document


load_dotenv()
client = OpenAI()

# Target size per chunk (bytes). 24MB leaves buffer under 25MB Whisper limit.
CHUNK_TARGET_SIZE = 24 * 1024 * 1024


def sanitize_for_fpdf(text):
    replacements = {
        "—": "-",    # em dash
        "–": "-",    # en dash
        "‐": "-",    # hyphen (U+2010)
        "―": "-",    # horizontal bar (U+2015)
        "−": "-",    # minus sign
        "‑": "-",    # non-breaking hyphen (U+2011)
        "“": '"',    # left double quotation mark
        "”": '"',    # right double quotation mark
        "„": '"',    # double low-9 quotation mark
        "‟": '"',    # double high-reversed-9 quotation mark
        "‘": "'",    # left single quotation mark
        "’": "'",    # right single quotation mark
        "‚": "'",    # single low-9 quotation mark
        "‛": "'",    # single high-reversed-9 quotation mark
        "…": "...",  # ellipsis
        "•": "-",    # bullet
        "‧": ".",    # hyphenation point
        "·": ".",    # middle dot
        " ": " ",    # narrow no-break space (U+202F)
        " ": " ",    # thin space (U+2009)
        "\u00A0": " ",  # no-break space
        "\u200B": "",   # zero-width space (remove)
        "\u200C": "",   # zero-width non-joiner (remove)
        "\u200D": "",   # zero-width joiner (remove)

        # Math symbols replacements
        "∫": "integral of ",
        "×": "x",    # multiplication sign → letter x
        "÷": "/",    # division sign → slash
        "√": "sqrt", # square root → textual substitute
        "α": "alpha",
        "β": "beta",
        "γ": "gamma",
        "Δ": "Delta",
        "∞": "infinity",
        "≈": "~",
        "≠": "!=",
        "≤": "<=",
        "≥": ">=",

    }

    for original, replacement in replacements.items():
        text = text.replace(original, replacement)

    # Finally encode to latin-1, replacing unsupported chars with '?'
    return text.encode("latin-1", errors="replace").decode("latin-1")

def split_audio_by_size(file_path, chunk_target_size=CHUNK_TARGET_SIZE):
    audio = AudioSegment.from_file(file_path)
    total_size = os.path.getsize(file_path)
    duration_ms = len(audio)

    # Estimate bytes per millisecond
    bytes_per_ms = total_size / duration_ms
    chunk_length_ms = math.floor(chunk_target_size / bytes_per_ms)

    chunks = []
    for i in range(0, duration_ms, chunk_length_ms):
        chunk = audio[i:i + chunk_length_ms]

        # Create a temp file name only, don't keep the file handle open
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
            chunk_path = tmp.name

        chunk.export(chunk_path, format="mp3")  # now safely write to it
        chunks.append(chunk_path)

    return chunks

def transcribe_audio(file_path):
    chunk_paths = split_audio_by_size(file_path)
    full_transcript = ""

    for i, chunk_path in enumerate(chunk_paths):
        with open(chunk_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )
            full_transcript += transcript.text + "\n\n"
        os.remove(chunk_path)

    return full_transcript.strip()

def format_transcription(text):
    chunks = chunk_text_by_tokens(text)
    formatted_chunks = []


    for i, chunk in enumerate(chunks):
        response = client.chat.completions.create(
            model="o4-mini-2025-04-16",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that formats audio transcripts."},
                {"role": "user", "content": f"Add punctuation and paragraphing to this transcript:\n{chunk}"}
            ],
            max_completion_tokens=40000
        )
        formatted_chunks.append(response.choices[0].message.content.strip())


    return "\n\n".join(formatted_chunks)

def summarise_text_from_transcript(text):
    def safe_request(prompt, context_name="summary"):
        try:
            response = client.chat.completions.create(
                model="o4-mini-2025-04-16",
                messages=prompt,
                max_completion_tokens=20000
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            return None

    # Use smaller chunk size to prevent overflow
    chunks = chunk_text_by_tokens(text, max_tokens=20000)
    partial_summaries = []

    # One-shot summary if small
    if len(chunks) == 1:
        prompt = [
            {"role": "system", "content": "You are a helpful assistant that summarizes transcripts."},
            {"role": "user", "content": f"Please summarize the following transcript into a few paragraphs. The first line should be a title (no more than 9 words):\n\n{chunks[0]}"}
        ]
        summary = safe_request(prompt, "one-shot summary")
        if not summary:
            return "[ERROR] Summary failed."
        return summary

    # Chunked summarization
    for i, chunk in enumerate(chunks):
        prompt = [
            {"role": "system", "content": "You are a helpful assistant that summarizes transcripts."},
            {"role": "user", "content": f"Summarize this part of a transcript into a paragraph:\n{chunk}"}
        ]
        summary = safe_request(prompt, f"chunk {i + 1}")
        if summary:
            partial_summaries.append(summary)
        else:
            partial_summaries.append(f"(Chunk {i+1} could not be summarized.)")


    combined = "\n\n".join(partial_summaries)

    # Trim if over safe token limit (leave room for output)
    enc = tiktoken.get_encoding("cl100k_base")
    combined_tokens = enc.encode(combined)
    MAX_FINAL_INPUT_TOKENS = 15000

    if len(combined_tokens) > MAX_FINAL_INPUT_TOKENS:

        combined = enc.decode(combined_tokens[:MAX_FINAL_INPUT_TOKENS])

    final_prompt = [
        {"role": "system", "content": "You are a helpful assistant that summarizes summaries."},
        {"role": "user", "content": f"Combine and refine the following summaries into a few concise paragraphs. The first line should be a title (no more than 9 words):\n\n{combined}"}
    ]

    final_summary = safe_request(final_prompt, "final summary")


    return final_summary or "[ERROR] Final summary could not be generated."

def chunk_text_by_tokens(text, max_tokens=20000):
    enc = tiktoken.get_encoding("cl100k_base")
    paragraphs = text.split("\n")

    chunks = []
    current_chunk = []
    current_tokens = 0

    for para in paragraphs:
        token_count = len(enc.encode(para))
        if token_count > max_tokens:
            # Paragraph too large: split by sentence or truncate
            sentences = para.split(". ")
            temp_chunk = []
            temp_tokens = 0
            for sentence in sentences:
                sentence_tokens = len(enc.encode(sentence))
                if temp_tokens + sentence_tokens > max_tokens:
                    chunks.append(". ".join(temp_chunk))
                    temp_chunk = [sentence]
                    temp_tokens = sentence_tokens
                else:
                    temp_chunk.append(sentence)
                    temp_tokens += sentence_tokens
            if temp_chunk:
                chunks.append(". ".join(temp_chunk))
        elif current_tokens + token_count > max_tokens:
            chunks.append("\n".join(current_chunk))
            current_chunk = [para]
            current_tokens = token_count
        else:
            current_chunk.append(para)
            current_tokens += token_count

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def generate_pdf_from_text(title, body_lines, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, sanitize_for_fpdf(title), ln=True, align="C")
    pdf.ln(10)

    # Body
    pdf.set_font("Arial", size=12)


    # Group lines into paragraphs by detecting blank lines
    paragraphs = []
    paragraph = []

    for line in body_lines:
        stripped = line.strip()
        if stripped:
            paragraph.append(sanitize_for_fpdf(stripped))
        elif paragraph:
            # Preserve line breaks within the paragraph
            paragraphs.append("\n".join(paragraph))
            paragraph = []
    if paragraph:
        paragraphs.append("\n".join(paragraph))

    # Write each paragraph with spacing
    for para in paragraphs:
        pdf.multi_cell(0, 6, para)
        pdf.ln(4)  # Add space between paragraphs

    pdf.output(output_path)


def generate_word_doc_from_text(title, body_lines, output_path):


    doc = Document()

    # Add title as heading
    doc.add_heading(title, level=1)


    for line in body_lines:
        doc.add_paragraph(line)

    doc.save(output_path)

def generate_latex_from_transcript(transcript_text, output_dir="uploads", tex_filename="transcript_body.tex"):
    import os

    os.makedirs(output_dir, exist_ok=True)
    tex_path = os.path.join(output_dir, tex_filename)

    chunks = chunk_text_by_tokens(transcript_text, max_tokens=20000)
    latex_bodies = []

    for i, chunk in enumerate(chunks):

        response = client.chat.completions.create(
            model="o4-mini-2025-04-16",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that converts transcripts to LaTeX."},
                {"role": "user", "content": f"Convert this into LaTeX body code. Escape all special characters where necessary. Do NOT include document preamble or \\begin{{document}}:\n\n{chunk}"}
            ],
            max_completion_tokens=40000
        )

        body = response.choices[0].message.content.strip()
        if body.startswith("```"):
            body = "\n".join(body.splitlines()[1:-1])

        clean_body = clean_latex_unicode(body)
        latex_bodies.append(clean_body)

    final_tex = (
        "\\documentclass{article}\n"
        "\\usepackage[margin=1in]{geometry}\n"
        "\\usepackage{amsmath, amssymb}\n"
        "\\usepackage{enumitem}\n"
        "\\usepackage{url}\n"
        "\\begin{document}\n\n"
        + "\n\n".join(latex_bodies)
        + "\n\n\\end{document}"
    )

    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(final_tex)


    return tex_path


def generate_latex_summary(transcript_text, output_dir="uploads", tex_filename="summary_body.tex "):
    import os

    os.makedirs(output_dir, exist_ok=True)
    tex_path = os.path.join(output_dir, tex_filename)

    chunks = chunk_text_by_tokens(transcript_text, max_tokens=20000)
    latex_bodies = []

    for i, chunk in enumerate(chunks):

        response = client.chat.completions.create(
            model="o4-mini-2025-04-16",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that converts transcripts to LaTeX summaries."},
                {"role": "user",
                 "content": f"Convert this into a summary in LaTeX body code. Escape all special characters where necessary. Do NOT include document preamble or \\begin{{document}}:\n\n{chunk}"}
            ],
            max_completion_tokens=40000
        )

        body = response.choices[0].message.content.strip()
        if body.startswith("```"):
            body = "\n".join(body.splitlines()[1:-1])

        clean_body = clean_latex_unicode(body)
        latex_bodies.append(clean_body)

    final_tex = (
            "\\documentclass{article}\n"
            "\\usepackage[margin=1in]{geometry}\n"
            "\\usepackage{amsmath, amssymb}\n"
            "\\usepackage{enumitem}\n"
            "\\usepackage{url}\n"
            "\\begin{document}\n\n"
            + "\n\n".join(latex_bodies)
            + "\n\n\\end{document}"
    )

    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(final_tex)

    return tex_path


def clean_latex_unicode(text):

    replacements = {
        # Dashes and quotes (text mode)
        "−": "-",    # minus
        "–": "-",    # en dash
        "—": "--",   # em dash
        "“": "``",
        "”": "''",
        "‘": "`",
        "’": "'",
        "‚": ",",
        "„": ",,",
        "…": "...",
        "•": r"\textbullet{}",

        # Fractions and symbols (text mode)
        "¼": r"\textonequarter{}",
        "½": r"\textonehalf{}",
        "¾": r"\textthreequarters{}",
        "©": r"\textcopyright{}",
        "®": r"\textregistered{}",
        "™": r"\texttrademark{}",
        "€": r"\euro{}",
        "£": r"\pounds{}",
        "°": r"$^\circ$",

        # Math operators (math mode)
        "×": r"$\times$",
        "÷": r"$\div$",
        "±": r"$\pm$",
        "∓": r"$\mp$",
        "≈": r"$\approx$",
        "≠": r"$\neq$",
        "≤": r"$\leq$",
        "≥": r"$\geq$",
        "∑": r"$\sum$",
        "∏": r"$\prod$",
        "√": r"$\sqrt{}$",
        "∞": r"$\infty$",
        "∫": r"$\int$",
        "∂": r"$\partial$",
        "∇": r"$\nabla$",
        "∈": r"$\in$",
        "∉": r"$\notin$",
        "∩": r"$\cap$",
        "∪": r"$\cup$",
        "⊂": r"$\subset$",
        "⊃": r"$\supset$",
        "⊆": r"$\subseteq$",
        "⊇": r"$\supseteq$",
        "∧": r"$\land$",
        "∨": r"$\lor$",
        "¬": r"$\neg$",
        "∀": r"$\forall$",
        "∃": r"$\exists$",
        "⇒": r"$\Rightarrow$",
        "⇐": r"$\Leftarrow$",
        "⇔": r"$\Leftrightarrow$",
        "→": r"$\rightarrow$",
        "←": r"$\leftarrow$",
        "↔": r"$\leftrightarrow$",

        # Greek lowercase (math mode)
        "α": r"$\alpha$",
        "β": r"$\beta$",
        "γ": r"$\gamma$",
        "δ": r"$\delta$",
        "ε": r"$\epsilon$",
        "ζ": r"$\zeta$",
        "η": r"$\eta$",
        "θ": r"$\theta$",
        "ι": r"$\iota$",
        "κ": r"$\kappa$",
        "λ": r"$\lambda$",
        "μ": r"$\mu$",
        "ν": r"$\nu$",
        "ξ": r"$\xi$",
        "ο": "o",  # not a math symbol
        "π": r"$\pi$",
        "ρ": r"$\rho$",
        "σ": r"$\sigma$",
        "τ": r"$\tau$",
        "υ": r"$\upsilon$",
        "φ": r"$\phi$",
        "χ": r"$\chi$",
        "ψ": r"$\psi$",
        "ω": r"$\omega$",

        # Greek uppercase (math mode)
        "Γ": r"$\Gamma$",
        "Δ": r"$\Delta$",
        "Θ": r"$\Theta$",
        "Λ": r"$\Lambda$",
        "Ξ": r"$\Xi$",
        "Π": r"$\Pi$",
        "Σ": r"$\Sigma$",
        "Υ": r"$\Upsilon$",
        "Φ": r"$\Phi$",
        "Ψ": r"$\Psi$",
        "Ω": r"$\Omega$",

        # Variant letters (math mode)
        "ϵ": r"$\varepsilon$",
        "ϑ": r"$\vartheta$",
        "ϕ": r"$\varphi$",
        "ς": r"$\varsigma$",

        # Arrows (math mode)
        "↦": r"$\mapsto$",
        "∘": r"$\circ$",
        "∙": r"$\cdot$",
        "↗": r"$\nearrow$",
        "↘": r"$\searrow$",
        "↙": r"$\swarrow$",
        "↖": r"$\nwarrow$",
        "⇑": r"$\Uparrow$",
        "⇓": r"$\Downarrow$",

        # Superscripts
        "¹": r"$^{1}$",
        "²": r"$^{2}$",
        "³": r"$^{3}$",

        # Accented Latin letters (text mode)
        "á": r"\'{a}",
        "é": r"\'{e}",
        "í": r"\'{i}",
        "ó": r"\'{o}",
        "ú": r"\'{u}",
        "ñ": r"\~{n}",
        "ü": r"\"{u}",
        "ç": r"\c{c}",

        # Misc text symbols
        "¶": r"\P",
        "§": r"\S",
        "†": r"\dagger",
        "‡": r"\ddagger",
        "‰": r"\permil",
        "′": r"'",
        "″": r"''",
        "‴": r"'''",
        "⁄": "/",

        # Whitespace and spacing
        "\u00A0": " ",
        "\u2009": r"\,",        # thin space
        "\u2002": r"\enspace",  # en space
        "\u2003": r"\quad",     # em space
        "\u2011": "-",          # non-breaking hyphen
    }

    for bad_char, replacement in replacements.items():
        text = text.replace(bad_char, replacement)

    return text


def compile_latex_to_pdf(latex_file, pdf_path):

    output_dir = os.path.dirname(latex_file)
    tex_filename = os.path.basename(latex_file)

    try:
        # Run pdflatex, suppress output
        subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", tex_filename],
            cwd=output_dir,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True
        )

        generated_pdf = os.path.join(output_dir, tex_filename.replace(".tex", ".pdf"))

        # If desired pdf_path is different, move the file
        if os.path.abspath(generated_pdf) != os.path.abspath(pdf_path):
            shutil.move(generated_pdf, pdf_path)


    except (subprocess.CalledProcessError, FileNotFoundError):
        pass




def generate_latex_pdf_from_transcipt(transcript, pdf_path):
    latex_file = generate_latex_from_transcript(transcript, "uploads", "Math_Transcription.tex")
    compile_latex_to_pdf(latex_file, pdf_path)

def generate_latex_pdf_from_summary(transcript, pdf_path):
    latex_file = generate_latex_from_transcript(transcript, "uploads", "Math_Summary.tex")
    compile_latex_to_pdf(latex_file, pdf_path)